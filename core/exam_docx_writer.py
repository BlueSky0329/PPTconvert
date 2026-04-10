"""将 ParsedExam 写入 Word：按「篇题 / 材料 / 题目 / 选项」输出。"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

from core.pdf_exam_models import ParsedExam, RichLine
from core.pdf_exam_parse import _normalize_digits


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def extract_year_region(section_title: str) -> tuple[str, str]:
    """从篇题（如 2026年·天津·资料分析）解析年份与地区。"""
    t = _normalize_digits((section_title or "").strip())
    year = ""
    m = re.search(r"(?<!\d)(\d{4})(?!\d)", t)
    if m:
        year = m.group(1)
    region = ""
    m2 = re.search(r"\d{4}\s*年\s*[·•．.\s]*\s*([^·•．\n]+?)\s*[·•．]", t)
    if m2:
        region = m2.group(1).strip()
    else:
        m3 = re.search(
            r"\d{4}\s*年\s*[·•．.\s]*\s*([^·•．\n]+?)\s*(?:资料分析|数量关系)",
            t,
        )
        if m3:
            region = m3.group(1).strip()
    return year, region


def _add_rich_paragraph(doc: Document, rich: RichLine, font_name: str, size: Pt, bold: bool = False):
    p = doc.add_paragraph()
    if not rich.parts:
        return p
    for text, img in rich.parts:
        if img and os.path.isfile(img):
            run = p.add_run()
            try:
                run.add_picture(img, width=Inches(5.8))
            except Exception:
                run = p.add_run(f"[图片:{os.path.basename(img)}]")
                _set_run_font(run, font_name)
                run.font.size = size
                run.bold = bold
        elif text:
            run = p.add_run(text)
            _set_run_font(run, font_name)
            run.font.size = size
            run.bold = bold
    return p


def _set_paragraph_runs_font(paragraph, font_name: str, size: Pt, bold: bool = False):
    for run in paragraph.runs:
        _set_run_font(run, font_name)
        run.font.size = size
        run.bold = bold


def _is_text_only_line(rich: RichLine) -> bool:
    return bool(rich.parts) and all(not img for _text, img in rich.parts)


def _coalesce_rich_lines(lines: list[RichLine]) -> list[RichLine]:
    merged: list[RichLine] = []
    buffer = ""

    def flush_text():
        nonlocal buffer
        if buffer:
            merged.append(RichLine(parts=[(buffer, None)]))
            buffer = ""

    for line in lines:
        if _is_text_only_line(line):
            text = "".join(text_part for text_part, _img in line.parts)
            if not text:
                continue
            buffer += text
            continue
        flush_text()
        merged.append(line)

    flush_text()
    return merged


def _prepend_question_number(lines: list[RichLine], source_number: str) -> list[RichLine]:
    if not source_number:
        return list(lines)
    numbered = list(lines)
    for idx, line in enumerate(numbered):
        if not _is_text_only_line(line):
            continue
        text = "".join(text_part for text_part, _img in line.parts)
        numbered[idx] = RichLine(parts=[(f"{source_number}. {text}", None)])
        return numbered
    return [RichLine(parts=[(f"{source_number}. ", None)])] + numbered


def write_parsed_exam_to_docx(
    exam: ParsedExam,
    out_path: str,
    font_name: str = "微软雅黑",
    stem_size_pt: int = 11,
    option_size_pt: int = 11,
) -> None:
    doc = Document()
    stem_size = Pt(stem_size_pt)
    opt_size = Pt(option_size_pt)

    def write_question_section(title: str, questions: list) -> None:
        p_sec = doc.add_paragraph(title.strip())
        _set_paragraph_runs_font(p_sec, font_name, stem_size, bold=True)
        for q in questions:
            stem_lines = _prepend_question_number(
                _coalesce_rich_lines(q.stem_lines),
                q.source_number,
            )
            for rl in stem_lines:
                _add_rich_paragraph(doc, rl, font_name, stem_size, bold=False)
            for rl in q.option_lines:
                _add_rich_paragraph(doc, rl, font_name, opt_size, bold=False)
            doc.add_paragraph("")

    for sec in exam.politics_sections:
        write_question_section(sec.title, sec.questions)

    for sec in exam.common_sense_sections:
        write_question_section(sec.title, sec.questions)

    for sec in exam.verbal_sections:
        write_question_section(sec.title, sec.questions)

    for sec in exam.data_sections:
        p_sec = doc.add_paragraph(sec.title.strip())
        _set_paragraph_runs_font(p_sec, font_name, stem_size, bold=True)
        for mu in sec.materials:
            p_header = doc.add_paragraph(mu.header.strip())
            _set_paragraph_runs_font(p_header, font_name, stem_size, bold=True)
            for rl in _coalesce_rich_lines(mu.intro_lines):
                _add_rich_paragraph(doc, rl, font_name, stem_size, bold=False)
            for qi in range(len(mu.questions)):
                q = mu.questions[qi]
                stem_lines = _prepend_question_number(
                    _coalesce_rich_lines(q.stem_lines),
                    q.source_number,
                )
                for rl in stem_lines:
                    _add_rich_paragraph(doc, rl, font_name, stem_size, bold=False)
                for rl in q.option_lines:
                    _add_rich_paragraph(doc, rl, font_name, opt_size, bold=False)
                if qi != len(mu.questions) - 1:
                    doc.add_paragraph("")
            doc.add_paragraph("")

    for sec in exam.quant_sections:
        write_question_section(sec.title, sec.questions)

    for sec in exam.reasoning_sections:
        write_question_section(sec.title, sec.questions)

    doc.save(out_path)


def parsed_exam_is_empty(exam: ParsedExam) -> bool:
    return (
        not exam.politics_sections
        and not exam.common_sense_sections
        and not exam.verbal_sections
        and not exam.data_sections
        and not exam.quant_sections
        and not exam.reasoning_sections
    )
