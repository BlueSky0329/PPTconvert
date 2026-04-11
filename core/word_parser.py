import logging
import re
from typing import Iterator, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from core.image_extractor import ImageExtractor
from core.models import Option, Question
from core.subject_inference import default_subject_title, infer_document_subject, infer_subject_from_content
from core.word_math import paragraph_full_text, paragraph_has_drawing

LOGGER = logging.getLogger(__name__)

SECTION_HEADER = re.compile(r"^\d+\s*[.\uFF0E\u3001]\s*[\u4e00-\u9fff]+$")
SECTION_PREFIX = re.compile(
    r"^\s*(?:[一二三四五六七八九十百千万\d]+|[IVXLC]+)\s*[.\uFF0E\u3001,:：]?\s*",
    re.IGNORECASE,
)
OPTION_MARKER = re.compile(
    r"(?<![A-Za-z])([A-Z])\s*[.\uFF0E\u3001)\uFF09]\s*",
    re.IGNORECASE,
)
OPTION_LINE = re.compile(r"^\s*[A-Z]\s*[.\uFF0E\u3001)\uFF09]\s*", re.IGNORECASE)
ANSWER_LINE = re.compile(
    r"(?:\u7b54\u6848|\u6b63\u786e\u7b54\u6848)\s*[\uff1a:]\s*([A-Za-z](?:\s*[,/\u3001\uff0c\s]\s*[A-Za-z])*)",
    re.IGNORECASE,
)
MATERIAL_HEADER = re.compile(r"^\s*材料[一二三四五六七八九十百千万\d〇零两]+\s*$")

QUESTION_START_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    (
        "year_area",
        re.compile(
            r"^(?P<label>[\uFF08(]\s*\d{4}(?:\s*[\u00b7.\uFF0E-]\s*[^)\uFF09]+)?\s*[)\uFF09])\s*(?P<stem>.*)$"
        ),
    ),
    (
        "chinese_number",
        re.compile(
            r"^\u7b2c\s*(?P<number>\d{1,3})\s*\u9898(?:\s*[\uff1a:.\uFF0E\u3001)\uFF09]\s*)?(?P<stem>.*)$"
        ),
    ),
    (
        "numeric_paren",
        re.compile(r"^[\uFF08(]\s*(?P<number>\d{1,3})\s*[)\uFF09]\s*(?P<stem>.*)$"),
    ),
    (
        "numeric_plain",
        re.compile(r"^(?P<number>\d{1,3})\s*[.\uFF0E\u3001)\uFF09]\s*(?P<stem>.*)$"),
    ),
)


def _paragraph_has_image(paragraph) -> bool:
    return paragraph_has_drawing(paragraph)


def _get_paragraph_text(paragraph) -> str:
    return paragraph_full_text(paragraph)


def iter_all_paragraphs(doc) -> Iterator[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in doc.tables:
        yield from walk_table(table)


def _match_question_start(text: str) -> Optional[tuple[Optional[str], str]]:
    stripped = text.strip()
    for pattern_name, pattern in QUESTION_START_PATTERNS:
        match = pattern.match(stripped)
        if not match:
            continue
        if pattern_name == "year_area":
            return match.group("label").strip(), (match.group("stem") or "").strip()
        return None, (match.groupdict().get("stem") or "").strip()
    return None


def _extract_question_number(text: str) -> Optional[str]:
    stripped = text.strip()
    for pattern_name, pattern in QUESTION_START_PATTERNS:
        match = pattern.match(stripped)
        if not match:
            continue
        if pattern_name == "year_area":
            return None
        number = match.groupdict().get("number")
        return number.strip() if number else None
    return None


def _normalize_answer(raw: str) -> Optional[str]:
    letters: list[str] = []
    for letter in re.findall(r"[A-Za-z]", raw.upper()):
        if letter not in letters:
            letters.append(letter)
    return "".join(letters) or None


def _parse_options_from_line(text: str) -> list[Option]:
    matches = list(OPTION_MARKER.finditer(text))
    if not matches:
        return []

    options: list[Option] = []
    for idx, match in enumerate(matches):
        letter = match.group(1).upper()
        body_start = match.end()
        body_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        option_text = text[body_start:body_end].strip()
        options.append(Option(letter=letter, text=option_text))
    return options


def _section_kind_from_text(text: str) -> Optional[str]:
    stripped = text.strip()
    if not stripped or _match_question_start(stripped) is not None:
        return None
    normalized = SECTION_PREFIX.sub("", stripped)
    if "资料分析" in normalized:
        return "data"
    if "数量关系" in normalized:
        return "quant"
    if "政治理论" in normalized:
        return "politics"
    if "常识判断" in normalized:
        return "common_sense"
    if "言语理解" in normalized:
        return "verbal"
    if "判断推理" in normalized:
        return "reasoning"
    if len(stripped) > 60 and not SECTION_HEADER.match(stripped):
        return None
    return None


def _material_header_from_text(text: str) -> Optional[str]:
    stripped = text.strip()
    if MATERIAL_HEADER.match(stripped):
        return stripped
    return None


class WordParser:
    def __init__(self, temp_dir: Optional[str] = None, document_subject_hint: Optional[str] = None):
        self.image_extractor = ImageExtractor(temp_dir)
        self.document_subject_hint = (document_subject_hint or "").strip().lower() or None

    @staticmethod
    def _build_material_prefixed_stem(
        stem: str,
        material_header: Optional[str],
        material_paragraphs: list[str],
    ) -> str:
        parts: list[str] = []
        if material_header:
            parts.append(material_header)
        parts.extend(p for p in material_paragraphs if p)
        if stem:
            parts.append(stem)
        return "\n".join(parts)

    def _append_images(self, question: Question, paragraph, question_number: int) -> None:
        image_paths = self.image_extractor.extract_from_paragraph(paragraph, question_number)
        question.image_paths.extend(image_paths)
        question.question_image_paths.extend(image_paths)

    def _new_question(
        self,
        q_counter: int,
        stem: str,
        source_label: Optional[str],
        source_question_number: Optional[str],
        section_kind: Optional[str],
        section_title: Optional[str],
        material_header: Optional[str],
        material_paragraphs: list[str],
        material_images: list[str],
        question_images: Optional[list[str]] = None,
    ) -> Question:
        material_text = "\n".join(p for p in material_paragraphs if p).strip() or None
        question_image_paths = list(question_images or [])
        material_image_paths = list(material_images)
        return Question(
            number=q_counter,
            stem=stem,
            source_label=source_label,
            source_question_number=source_question_number,
            material_header=material_header,
            material_text=material_text,
            image_paths=material_image_paths + question_image_paths,
            question_image_paths=question_image_paths,
            material_image_paths=material_image_paths,
            section_kind=section_kind,
            section_title=section_title,
        )

    def _finalize_question(
        self,
        questions: list[Question],
        current_question: Optional[Question],
    ) -> Optional[Question]:
        if current_question is None:
            return None

        if current_question.is_complete:
            questions.append(current_question)
            return None

        if current_question.display_stem or current_question.options or current_question.image_paths:
            LOGGER.warning(
                "Skipping incomplete question %s with %s option(s)",
                current_question.number,
                len(current_question.options),
            )
        return None

    def parse(self, docx_path: str) -> list[Question]:
        doc = Document(docx_path)
        paragraphs = list(iter_all_paragraphs(doc))
        questions: list[Question] = []
        current_question: Optional[Question] = None
        q_counter = 0
        current_section_kind: Optional[str] = None
        current_section_title: Optional[str] = None
        current_material_header: Optional[str] = None
        current_material_paragraphs: list[str] = []
        current_material_images: list[str] = []
        pending_data_texts: list[str] = []
        pending_data_images: list[str] = []

        document_hint = self.document_subject_hint
        if document_hint is None:
            raw_lines = [_get_paragraph_text(paragraph) for paragraph in paragraphs]
            image_count = sum(1 for paragraph in paragraphs if _paragraph_has_image(paragraph))
            material_header_count = sum(1 for line in raw_lines if _material_header_from_text(line or ""))
            inferred_kind, confidence = infer_document_subject(
                raw_lines,
                image_count=image_count,
                material_header_count=material_header_count,
            )
            if inferred_kind and confidence >= 0.55:
                document_hint = inferred_kind

        if document_hint:
            current_section_kind = document_hint
            current_section_title = default_subject_title(document_hint) if document_hint != "unknown" else None

        for paragraph in paragraphs:
            text = _get_paragraph_text(paragraph).strip()
            has_img = _paragraph_has_image(paragraph)

            if not text and not has_img:
                continue

            section_kind = _section_kind_from_text(text) if text else None
            if section_kind:
                current_question = self._finalize_question(questions, current_question)
                current_section_kind = section_kind
                current_section_title = text
                current_material_header = None
                current_material_paragraphs = []
                current_material_images = []
                pending_data_texts = []
                pending_data_images = []
                continue

            material_header = _material_header_from_text(text) if text else None
            if current_section_kind != "data" and material_header:
                current_question = self._finalize_question(questions, current_question)
                current_section_kind = "data"
                current_section_title = default_subject_title("data")
                current_material_header = material_header
                current_material_paragraphs = []
                current_material_images = []
                pending_data_texts = []
                pending_data_images = []
                continue

            if current_section_kind == "data" and material_header:
                current_question = self._finalize_question(questions, current_question)
                current_material_header = material_header
                current_material_paragraphs = []
                current_material_images = []
                pending_data_texts = []
                pending_data_images = []
                continue

            matched_question = _match_question_start(text) if text else None
            if text and SECTION_HEADER.match(text) and matched_question is None:
                continue

            if text and current_question:
                answer_match = ANSWER_LINE.search(text)
                if answer_match:
                    normalized_answer = _normalize_answer(answer_match.group(1))
                    if normalized_answer:
                        current_question.answer = normalized_answer
                    continue

            if matched_question is not None:
                source_label, stem = matched_question
                source_question_number = _extract_question_number(text) if text else None
                current_question = self._finalize_question(questions, current_question)
                q_counter += 1
                if current_section_kind == "data":
                    if pending_data_texts or pending_data_images:
                        current_material_paragraphs = list(pending_data_texts)
                        current_material_images = list(pending_data_images)
                        pending_data_texts = []
                        pending_data_images = []
                    current_question = self._new_question(
                        q_counter,
                        stem,
                        source_label,
                        source_question_number,
                        current_section_kind,
                        current_section_title,
                        current_material_header,
                        current_material_paragraphs,
                        current_material_images,
                    )
                else:
                    current_question = Question(
                        number=q_counter,
                        stem=stem,
                        source_label=source_label,
                        source_question_number=source_question_number,
                        section_kind=current_section_kind,
                        section_title=current_section_title,
                    )
                if has_img:
                    self._append_images(current_question, paragraph, q_counter)
                continue

            if text and current_question and OPTION_LINE.match(text):
                options = _parse_options_from_line(text)
                if not options:
                    LOGGER.warning(
                        "Question %s looks like an option line but could not be parsed: %r",
                        current_question.number,
                        text,
                    )
                else:
                    current_question.options.extend(options)
                if has_img:
                    self._append_images(current_question, paragraph, q_counter)
                continue

            if text and current_question is None and current_section_kind == "data" and OPTION_LINE.match(text):
                options = _parse_options_from_line(text)
                if options:
                    q_counter += 1
                    material_context_established = bool(current_material_paragraphs or current_material_images)
                    if material_context_established:
                        stem = "\n".join(pending_data_texts).strip()
                        question_images = list(pending_data_images)
                    else:
                        if not pending_data_texts:
                            LOGGER.warning("Skipping option-only block in data section: %r", text)
                            continue
                        stem = pending_data_texts[-1].strip()
                        current_material_paragraphs = list(pending_data_texts[:-1])
                        current_material_images = list(pending_data_images)
                        question_images = []
                    pending_data_texts = []
                    pending_data_images = []
                    current_question = self._new_question(
                        q_counter,
                        stem,
                        None,
                        None,
                        current_section_kind,
                        current_section_title,
                        current_material_header,
                        current_material_paragraphs,
                        current_material_images,
                        question_images=question_images,
                    )
                    current_question.options.extend(options)
                    if has_img:
                        self._append_images(current_question, paragraph, q_counter)
                    continue

            if has_img and current_question:
                self._append_images(current_question, paragraph, q_counter)
                continue

            if current_section_kind == "data" and current_question is None:
                if text:
                    pending_data_texts.append(text)
                if has_img:
                    pending_data_images.extend(
                        self.image_extractor.extract_from_paragraph(paragraph, q_counter + 1)
                    )
                continue

            if (
                current_section_kind == "data"
                and current_question
                and current_question.is_complete
                and text
                and not OPTION_LINE.match(text)
            ):
                current_question = self._finalize_question(questions, current_question)
                pending_data_texts = [text]
                pending_data_images = []
                if has_img:
                    pending_data_images.extend(
                        self.image_extractor.extract_from_paragraph(paragraph, q_counter + 1)
                    )
                continue

            if current_question and text:
                if current_question.options:
                    last_option = current_question.options[-1]
                    last_option.text = " ".join(
                        part for part in (last_option.text, text) if part
                    )
                else:
                    current_question.stem = "\n".join(
                        part for part in (current_question.stem, text) if part
                    )

        self._finalize_question(questions, current_question)
        inferred_pairs: list[tuple[str, float]] = []
        for question in questions:
            inferred_kind, confidence = infer_subject_from_content(
                stem=question.stem or "",
                options=[option.text for option in question.options],
                material_text=question.material_text or "",
                image_count=len(question.image_paths),
                material_header=question.material_header or "",
                allow_data=bool(question.material_header or question.material_text),
            )
            if inferred_kind == "unknown" and document_hint:
                inferred_kind = document_hint  # type: ignore[assignment]
            inferred_pairs.append((inferred_kind, confidence))

        for index in range(1, len(inferred_pairs) - 1):
            prev_kind = inferred_pairs[index - 1][0]
            current_kind, current_confidence = inferred_pairs[index]
            next_kind = inferred_pairs[index + 1][0]
            if current_kind != prev_kind and prev_kind == next_kind and current_confidence < 1.4:
                inferred_pairs[index] = (prev_kind, current_confidence)

        for question, (inferred_kind, confidence) in zip(questions, inferred_pairs):
            base_kind = (question.section_kind or "").strip().lower()
            strong_text_signal = len((question.stem or "").strip()) >= 10 or sum(
                len((option.text or "").strip()) for option in question.options
            ) >= 18
            if question.material_header or question.material_text:
                question.section_kind = "data"
            elif base_kind in {"", "unknown"}:
                question.section_kind = inferred_kind if inferred_kind != "unknown" else "unknown"
            elif inferred_kind not in {"unknown", base_kind} and confidence >= 0.9 and strong_text_signal:
                question.section_kind = inferred_kind
            else:
                question.section_kind = base_kind
            if not question.section_title or question.section_title == default_subject_title(base_kind or "unknown"):
                question.section_title = default_subject_title(question.section_kind)
        return questions

    def cleanup(self):
        self.image_extractor.cleanup()

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        self.cleanup()
