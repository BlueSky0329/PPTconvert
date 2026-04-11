import os
import tempfile
import unittest

from docx import Document

from core.word_parser import (
    WordParser,
    _extract_question_number,
    _material_header_from_text,
    _match_question_start,
    _normalize_answer,
    _parse_options_from_line,
    _section_kind_from_text,
)


class WordParserHelpersTest(unittest.TestCase):
    def test_match_question_start_supports_multiple_formats(self):
        self.assertEqual(
            _match_question_start("\uff082020\u00b7\u4e0a\u6d77\uff09\u9898\u5e72"),
            ("\uff082020\u00b7\u4e0a\u6d77\uff09", "\u9898\u5e72"),
        )
        self.assertEqual(
            _match_question_start("1. simple stem"),
            (None, "simple stem"),
        )
        self.assertEqual(
            _match_question_start("2\u3001another stem"),
            (None, "another stem"),
        )
        self.assertEqual(
            _match_question_start("\u7b2c3\u9898 \u4e2d\u6587\u9898\u5e72"),
            (None, "\u4e2d\u6587\u9898\u5e72"),
        )
        self.assertIsNone(_match_question_start("2024. not a supported plain question prefix"))

    def test_parse_options_from_line_supports_multiple_markers(self):
        options = _parse_options_from_line(
            "A. alpha B. beta C. gamma E. epsilon"
        )
        self.assertEqual([option.letter for option in options], ["A", "B", "C", "E"])
        self.assertEqual(options[-1].text, "epsilon")

    def test_normalize_answer_deduplicates_letters(self):
        self.assertEqual(_normalize_answer("A, C / A"), "AC")

    def test_material_and_section_detection(self):
        self.assertEqual(_section_kind_from_text("2026年·天津·资料分析"), "data")
        self.assertEqual(_section_kind_from_text("四. 数量关系:"), "quant")
        self.assertEqual(_section_kind_from_text("一. 政治理论：根据题目要求作答"), "politics")
        self.assertEqual(_section_kind_from_text("二、常识判断"), "common_sense")
        self.assertEqual(_section_kind_from_text("三. 言语理解与表达：本部分包括表达与理解"), "verbal")
        self.assertEqual(_section_kind_from_text("五. 判断推理：在四个选项中选出答案"), "reasoning")
        self.assertEqual(_material_header_from_text("材料一"), "材料一")
        self.assertIsNone(_material_header_from_text("题干正文"))
        self.assertEqual(_extract_question_number("12. 题干"), "12")


class WordParserIntegrationTest(unittest.TestCase):
    def test_parse_docx_with_numeric_and_chinese_question_prefixes(self):
        document = Document()
        document.add_paragraph("1. First stem")
        document.add_paragraph("A. Alpha")
        document.add_paragraph("B. Beta")
        document.add_paragraph("\u7b2c2\u9898 Second stem")
        document.add_paragraph("A. One")
        document.add_paragraph("B. Two")
        document.add_paragraph("\uff082020\u00b7\u4e0a\u6d77\uff09Third stem")
        document.add_paragraph("A. Three")
        document.add_paragraph("B. Four")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 3)
            self.assertEqual(questions[0].stem, "First stem")
            self.assertEqual(questions[1].stem, "Second stem")
            self.assertEqual(questions[2].stem, "Third stem")
            self.assertEqual(questions[2].source_label, "\uff082020\u00b7\u4e0a\u6d77\uff09")
            self.assertEqual([opt.letter for opt in questions[0].options], ["A", "B"])
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_data_material_is_prefixed_to_each_question(self):
        document = Document()
        document.add_paragraph("2026年·天津·资料分析")
        document.add_paragraph("材料一")
        document.add_paragraph("这是材料第一段。")
        document.add_paragraph("这是材料第二段。")
        document.add_paragraph("111. 第一题")
        document.add_paragraph("A. Alpha")
        document.add_paragraph("B. Beta")
        document.add_paragraph("112. 第二题")
        document.add_paragraph("A. Gamma")
        document.add_paragraph("B. Delta")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 2)
            self.assertEqual(questions[0].stem, "第一题")
            self.assertEqual(questions[1].stem, "第二题")
            self.assertEqual(questions[0].source_question_number, "111")
            self.assertEqual(questions[1].source_question_number, "112")
            self.assertEqual(questions[0].material_header, "材料一")
            self.assertEqual(questions[0].material_text, "这是材料第一段。\n这是材料第二段。")
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_data_material_without_question_numbers_can_still_parse(self):
        document = Document()
        document.add_paragraph("2026年·天津·资料分析")
        document.add_paragraph("材料一")
        document.add_paragraph("这是材料正文。")
        document.add_paragraph("第一道题题干")
        document.add_paragraph("A. Alpha")
        document.add_paragraph("B. Beta")
        document.add_paragraph("第二道题题干")
        document.add_paragraph("A. Gamma")
        document.add_paragraph("B. Delta")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 2)
            self.assertEqual(questions[0].stem, "第一道题题干")
            self.assertEqual(questions[1].stem, "第二道题题干")
            self.assertEqual(questions[0].material_header, "材料一")
            self.assertEqual(questions[0].material_text, "这是材料正文。")
            self.assertEqual([opt.letter for opt in questions[0].options], ["A", "B"])
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_parse_docx_without_titles_can_infer_quant_subject(self):
        document = Document()
        document.add_paragraph("甲、乙两车同时从两地相向而行，全程240千米，速度比为3:2，几小时后相遇？")
        document.add_paragraph("66. 第一题")
        document.add_paragraph("A. 4")
        document.add_paragraph("B. 5")
        document.add_paragraph("C. 6")
        document.add_paragraph("D. 8")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 1)
            self.assertEqual(questions[0].section_kind, "quant")
            self.assertEqual(questions[0].section_title, "数量关系")
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_parse_docx_without_titles_can_force_data_subject(self):
        document = Document()
        document.add_paragraph("材料一")
        document.add_paragraph("2024年某市工业增加值同比增长8.3%。")
        document.add_paragraph("111. 第一题")
        document.add_paragraph("A. 甲")
        document.add_paragraph("B. 乙")

        tmp_path = None
        parser = WordParser(document_subject_hint="data")
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 1)
            self.assertEqual(questions[0].section_kind, "data")
            self.assertEqual(questions[0].material_header, "材料一")
            self.assertIn("同比增长", questions[0].material_text or "")
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_parse_docx_can_reclassify_later_questions_when_title_missing(self):
        document = Document()
        document.add_paragraph("四. 数量关系")
        document.add_paragraph("66. 甲、乙两队合修一段公路，若甲单独修需要12天，乙单独修需要18天，两队合修几天完成？")
        document.add_paragraph("A. 6")
        document.add_paragraph("B. 7")
        document.add_paragraph("C. 8")
        document.add_paragraph("D. 9")
        document.add_paragraph("76. 如果所有甲都是乙，且有些乙是丙，那么下列哪项一定为真？")
        document.add_paragraph("A. 有些甲是丙")
        document.add_paragraph("B. 有些丙是甲")
        document.add_paragraph("C. 有些乙不是丙")
        document.add_paragraph("D. 所有甲都是乙")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 2)
            self.assertEqual(questions[0].section_kind, "quant")
            self.assertEqual(questions[1].section_kind, "reasoning")
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_parse_docx_can_switch_to_data_when_material_header_appears(self):
        document = Document()
        document.add_paragraph("四. 数量关系")
        document.add_paragraph("66. 甲、乙两地相距240千米，两车相向而行几小时后相遇？")
        document.add_paragraph("A. 4")
        document.add_paragraph("B. 5")
        document.add_paragraph("C. 6")
        document.add_paragraph("D. 8")
        document.add_paragraph("材料一")
        document.add_paragraph("2024年某市工业增加值同比增长8.3%。")
        document.add_paragraph("111. 根据上述材料，下列说法正确的是：")
        document.add_paragraph("A. 甲")
        document.add_paragraph("B. 乙")

        tmp_path = None
        parser = WordParser()
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            document.save(tmp_path)

            questions = parser.parse(tmp_path)
            self.assertEqual(len(questions), 2)
            self.assertEqual(questions[0].section_kind, "quant")
            self.assertEqual(questions[1].section_kind, "data")
            self.assertEqual(questions[1].material_header, "材料一")
        finally:
            parser.cleanup()
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
