import base64
import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

from domain.models import ExamProject, MaterialSet, PageRegion, PaperSource, QuestionNode, Section
from exporters.docx_booklet import export_project_to_docx

_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+yF9kAAAAASUVORK5CYII="
)


class DocxBookletTest(unittest.TestCase):
    def _write_png(self, path: str) -> None:
        with open(path, "wb") as file_obj:
            file_obj.write(_PNG_BYTES)

    def test_data_material_prefers_region_crop_over_flat_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            crop_path = os.path.join(temp_dir, "material.png")
            out_path = os.path.join(temp_dir, "output.docx")
            self._write_png(crop_path)

            project = ExamProject(
                title="test",
                source=PaperSource(pdf_path="sample.pdf"),
                sections=[
                    Section(
                        kind="data",
                        title="资料分析",
                        material_sets=[
                            MaterialSet(
                                material_id="data-1-1",
                                header="材料一",
                                body="2025年表格材料",
                                body_lines=["2025年表格材料"],
                                body_regions=[PageRegion(page_number=1, x0=10, y0=10, x1=100, y1=100)],
                                questions=[
                                    QuestionNode(
                                        source_number="111",
                                        stem="题干",
                                        options=[],
                                    )
                                ],
                            )
                        ],
                    )
                ],
            )

            with patch("exporters.docx_booklet.crop_material_regions", return_value=[crop_path]):
                export_project_to_docx(project, out_path)

            document = Document(out_path)
            paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn("资料分析", paragraph_texts)
            self.assertIn("材料一", paragraph_texts)
            self.assertIn("111. 题干", paragraph_texts)
            self.assertNotIn("2025年表格材料", paragraph_texts)
            self.assertEqual(len(document.inline_shapes), 1)

    def test_data_material_falls_back_to_flat_text_when_crop_missing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            out_path = os.path.join(temp_dir, "output.docx")
            project = ExamProject(
                title="test",
                source=PaperSource(pdf_path="sample.pdf"),
                sections=[
                    Section(
                        kind="data",
                        title="资料分析",
                        material_sets=[
                            MaterialSet(
                                material_id="data-1-1",
                                header="材料一",
                                body="2025年表格材料",
                                body_lines=["2025年表格材料"],
                                body_regions=[PageRegion(page_number=1, x0=10, y0=10, x1=100, y1=100)],
                                questions=[
                                    QuestionNode(
                                        source_number="111",
                                        stem="题干",
                                        options=[],
                                    )
                                ],
                            )
                        ],
                    )
                ],
            )

            with patch("exporters.docx_booklet.crop_material_regions", return_value=[]):
                export_project_to_docx(project, out_path)

            document = Document(out_path)
            paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn("2025年表格材料", paragraph_texts)
            self.assertEqual(len(document.inline_shapes), 0)


if __name__ == "__main__":
    unittest.main()
