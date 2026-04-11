from __future__ import annotations

import os
import shutil
import tempfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from domain.models import AssetRef, ExamProject, MaterialSet, QuestionNode
from exporters.material_crops import crop_material_regions


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.bold = bold


def _add_text_paragraph(doc: Document, text: str, font_name: str, size: Pt, bold: bool = False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, font_name, size, bold=bold)
    return paragraph


def _add_asset_paragraph(doc: Document, asset: AssetRef, width: float = 5.8):
    if not asset.path or not os.path.isfile(asset.path):
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(asset.path, width=Inches(width))


def _add_option_block(
    doc: Document,
    option,
    font_name: str,
    option_size: Pt,
    image_width: float = 3.4,
):
    paragraph = doc.add_paragraph()
    label_text = f"{option.letter}." if option.image_path and not option.text else f"{option.letter}. "

    label_run = paragraph.add_run(label_text)
    _set_run_font(label_run, font_name, option_size, bold=False)

    if option.text:
        body_run = paragraph.add_run(option.text)
        _set_run_font(body_run, font_name, option_size, bold=False)

    if option.image_path and os.path.isfile(option.image_path):
        image_run = paragraph.add_run()
        image_run.add_break()
        image_run.add_picture(option.image_path, width=Inches(image_width))

    return paragraph


def _render_material_body(
    doc: Document,
    project: ExamProject,
    material: MaterialSet,
    font_name: str,
    size: Pt,
    crop_dir: str | None,
) -> None:
    crop_paths: list[str] = []
    if crop_dir:
        crop_paths = crop_material_regions(project.source.pdf_path or "", material, crop_dir)
    if crop_paths:
        for path in crop_paths:
            _add_asset_paragraph(doc, AssetRef(kind="material_region_crop", path=path), width=6.0)
        return
    for line in material.body_lines:
        _add_text_paragraph(doc, line, font_name, size)
    for asset in material.body_assets:
        _add_asset_paragraph(doc, asset)


def _add_question_block(doc: Document, question: QuestionNode, font_name: str, stem_size: Pt, option_size: Pt):
    stem_prefix = f"{question.source_number}. " if question.source_number else ""
    _add_text_paragraph(doc, f"{stem_prefix}{question.stem}".strip(), font_name, stem_size)
    for asset in question.stem_assets:
        _add_asset_paragraph(doc, asset)
    for option in question.options:
        _add_option_block(doc, option, font_name, option_size)
    doc.add_paragraph("")


def export_project_to_docx(
    project: ExamProject,
    out_path: str,
    *,
    font_name: str = "微软雅黑",
    stem_size_pt: int = 11,
    option_size_pt: int = 11,
) -> str:
    document = Document()
    stem_size = Pt(stem_size_pt)
    option_size = Pt(option_size_pt)
    crop_dir = tempfile.mkdtemp(prefix="pptconvert_docx_material_crops_")
    try:
        for section in project.sections:
            if section.title:
                _add_text_paragraph(document, section.title.strip(), font_name, stem_size, bold=True)
            if section.kind == "data":
                for material in section.material_sets:
                    if not material.questions:
                        continue
                    _add_text_paragraph(document, material.header.strip(), font_name, stem_size, bold=True)
                    _render_material_body(
                        document,
                        project,
                        material,
                        font_name,
                        stem_size,
                        crop_dir,
                    )
                    for question in material.questions:
                        _add_question_block(document, question, font_name, stem_size, option_size)
            else:
                for question in section.questions:
                    _add_question_block(document, question, font_name, stem_size, option_size)

        output_dir = os.path.dirname(os.path.abspath(out_path))
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        document.save(out_path)
        return out_path
    finally:
        shutil.rmtree(crop_dir, ignore_errors=True)
